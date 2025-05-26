from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
import os
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import numpy as np
from datetime import datetime
from quchemreport.utils.units import nm_to_wnb

from cclib.parser.utils import PeriodicTable

def add_section_title(doc, title_text):
    section = doc.sections[0]
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), '009080')
    p._element.get_or_add_pPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'bottom', 'start', 'end', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = qn('w:' + edge)
            element = tcBorders.find(tag)
            if element is None:
                element = OxmlElement('w:' + edge)
                tcBorders.append(element)
            for key in edge_data:
                element.set(qn('w:' + key), str(edge_data[key]))

def display_vertical_lines(table, centralLines=True):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if centralLines:
                set_cell_border(
                    cell,
                    start={"val": "single", "sz": 2, "color": "000000", "space" : 0},
                    end={"val": "single", "sz": 2, "color": "000000", "space" : 0}
                )
            elif i == 0:
                set_cell_border(cell,start={"val": "single", "sz": 2, "color": "000000", "space": 0})
            elif i == len(row.cells) - 1:
                set_cell_border(cell,end={"val": "single", "sz": 2, "color": "000000", "space" : 0})

def set_vertical_line(table, index, visible):
    color = "000000" if visible else "FFFFFF"
    for row in table.rows:
        if index == len(row.cells):
            set_cell_border(row.cells[index-1], end={"val": "single", "sz": 2, "color": color, "space": 0})
        else :
            set_cell_border(row.cells[index], start={"val": "single", "sz": 2, "color": color, "space": 0})

def set_horizontal_line(table, index, visible=True, size=2):
    color = "000000" if visible else "FFFFFF"
    if index == len(table.rows):
        for cell in table.rows[index-1].cells:
            set_cell_border(cell, bottom={"val": "single", "sz": size, "color": color, "space": 0})
    else :
        for cell in table.rows[index].cells:
            set_cell_border(cell, top={"val": "single", "sz": size, "color": color, "space": 0})

def set_all_cell_borders(table, visible) :
    color = "000000" if visible else "FFFFFF"
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": 0, "color": color},
                bottom={"val": "single", "sz": 0, "color": color},
                start={"val": "single", "sz": 0, "color": color},
                end={"val": "single", "sz": 0, "color": color}
            )

def create_table(doc, t):
    table = doc.add_table(rows=len(t), cols=len(t[0]))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    table.style = 'Table Grid'
    for i in range(len(t)):
        row = table.rows[i]
        for j in range(len(t[i])):
            row.cells[j].text = str(t[i][j])
    set_all_cell_borders(table, False)
    tbl_pr = table._element.xpath('.//w:tblPr')[0]
    tbl_ind = OxmlElement('w:tblInd')
    tbl_ind.set(qn('w:w'), str(int(600)))
    tbl_ind.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_ind)
    return table

def figure_one_col(doc, img_path1, caption_text, width_in_inches=4):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run1 = p_img.add_run()
    run1.add_picture(img_path1, width=Inches(width_in_inches))
    doc.add_paragraph()
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption.runs[0]
    run.italic = True
    run.font.size = Pt(10)


def figure_two_col(doc, img_path1, img_path2, caption_text, width_in_inches=3.5):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    p1 = table.cell(0, 0).paragraphs[0]
    p1.add_run().add_picture(img_path1, width=Inches(width_in_inches))
    p1.paragraph_format.left_indent = None
    p1.paragraph_format.right_indent = None
    p1.paragraph_format.space_before = 0
    p1.paragraph_format.space_after = 0
    p1.style = doc.styles['Normal']
    table.cell(0, 1).paragraphs[0].add_run().add_picture(img_path2, width=Inches(width_in_inches))
    doc.add_paragraph()
    caption = doc.add_paragraph(caption_text)
    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = caption.runs[0]
    run.italic = True
    run.font.size = Pt(10)
    set_all_cell_borders(table, False)
    return table

def add_footer(doc):
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    p_bdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), 'auto')
    p._element.get_or_add_pPr().append(p_bdr)
    p_bdr.append(top)
    para = footer.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.left_indent = Pt(0)
    para.paragraph_format.right_indent = Pt(0)
    para.paragraph_format.tab_stops.add_tab_stop(Pt(500))
    now = datetime.now()
    day = now.day
    suffix = 'th' if 11 <= day <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    date_str = now.strftime(f"%A {day}{suffix} %B, %Y  %H:%M")
    run_left = para.add_run(date_str)
    run_left.font.size = Pt(10)
    run_right = para.add_run("\tPage ")
    run_right.font.size = Pt(10)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_right._r.append(fldChar1)
    run_right._r.append(instrText)
    run_right._r.append(fldChar2)
    """
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=1, width=1)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.allow_autofit = True
    cell = table.cell(0, 0)

    para = cell.paragraphs[0]
    para.style = doc.styles['Normal']

    # Définir tabulation à droite
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.SPACES)

    # Texte de gauche (date + heure)
    now = datetime.now().strftime("%A %dᵗʰ %B, %Y  %H:%M")
    run_left = para.add_run(now)
    run_left.font.size = Pt(10)

    # Texte de droite (après tabulation)
    run_right = para.add_run("\tPage ")
    run_right.font.size = Pt(10)

    # Champ dynamique PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_right._r.append(fldChar1)
    run_right._r.append(instrText)
    run_right._r.append(fldChar2)


def json2docx(config, json_list, data, mode="clean"):
    ### SECTION 1
    report_type = config.output.include.electron_density_difference.mode
    data_ref = data['data_for_discretization']
    job_types = data['job_types']
    name = data_ref["molecule"]["formula"]
    dirname = os.path.basename(os.getcwd())
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("MOLECULAR CALCULATION REPORT")
    run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    section_title = doc.add_paragraph()
    add_section_title(doc, "1. MOLECULE")
    run.bold = True
    run.font.size = Pt(14)
    section_title_format = section_title.paragraph_format
    section_title_format.space_after = Pt(6)
    figure_two_col(doc, "temp/img-TOPOLOGY.png", "temp/img-TOPOLOGY_cam2.png", "Figure 1: Chemical structure diagram with atomic numbering from two points of view.", 3)
    inchi = (data_ref["molecule"]["inchi"]).rstrip().split("=")[-1]
    t = [
        ["Directory name", dirname],
        ["Formula", data_ref["molecule"]["formula"]],
        ["Charge", data_ref["molecule"]["charge"]],
        ["Spin multiplicity", data_ref["molecule"]["multiplicity"]]]
    if report_type == 'full':
        t.append(["Monoisotopic mass", "%.5f Da" % data_ref["molecule"]["monoisotopic_mass"]])
        t.append(["InChI", inchi])
        if (len(data_ref["molecule"]["smi"]) < 80):
            t.append(["SMILES", data_ref["molecule"]["smi"]])
    table = create_table(doc, t)
    table.columns[0].width = Pt(160);
    table.columns[1].width = Pt(320);
    display_vertical_lines(table, False)
    doc.add_paragraph()


    ### SECTION 2
    add_section_title(doc, "2. COMPUTATIONAL DETAILS")
    t = []
    software = data_ref["comp_details"]["general"]["package"]
    try:
        t.append(['Software', data_ref["comp_details"]["general"]["package"], '(' + data_ref["comp_details"]["general"]["package_version"] + ')'])
    except KeyError:
        pass
    t.append(['Computational method', data_ref["comp_details"]["general"]["last_theory"], " "])
    t.append(['Functional', data_ref["comp_details"]["general"]["functional"], " "])
    try:
        t.append(['Basis set name', data_ref["comp_details"]["general"]["basis_set_name"], " "])
    except KeyError:
        pass
    t.append(['Number of basis set functions', data_ref["comp_details"]["general"]["basis_set_size"], " "])
    t.append(['Closed shell calculation', data_ref["comp_details"]["general"]["is_closed_shell"], " "])
    try:
        t.append(['Integration grid', data_ref["comp_details"]["general"]["integration_grid"], " "])
    except KeyError:
        pass
    try:
        t.append(['Solvent', data_ref["comp_details"]["general"]["solvent"], " "])
    except KeyError:
        pass
    #            TODO
    #            add_row_filter(param_table, ['Solvent reaction filed method',
    #                                         json_list[i]["comp_details"]["general"]["solvent_reaction_field"]])

    scfTargets = data_ref["comp_details"]["general"]["scf_targets"][-1]
    if software == "Gaussian":  # Gaussian or GAUSSIAN (upper/lower?
        t.append(["Requested SCF convergence on RMS and Max density matrix", scfTargets[0], scfTargets[1]])
        t.append(["Requested SCF convergence on energy", scfTargets[2], " "])
    if software == "GAMESS":
        t.append(["Requested SCF convergence on density", scfTargets[0], " "])

    # Specific calculations parameters :
    OPT_param_print = False
    for i, jsonfile in enumerate(json_list):
        # OPT calculation parameters :
        if (('OPT' in job_types[i]) or ('FREQ' in job_types[i] and 'OPT' in job_types[i])) and (OPT_param_print == False):
            t.append([" ", " ", " "])
            k = 0
            j = str(k + 1)
            try:
                t.append(["Job type: Geometry optimization", " ", " "])
                geomTargets = json_list[i]["comp_details"]["geometry"]["geometric_targets"]
                geomValues = json_list[i]["results"]['geometry']['geometric_values'][-1]
                if software == "Gaussian":  # Gaussian or GAUSSIAN (upper/lower?
                    t.append( ["Max Force value and threshold", "%.6f" % geomValues[0], "%.6f" % geomTargets[0]])
                    t.append(["RMS Force value and threshold", "%.6f" % geomValues[1], "%.6f" % geomTargets[1]])
                    t.append(["Max Displacement value and threshold", "%.6f" % geomValues[2], "%.6f" % geomTargets[2]])
                    t.append(["RMS Displacement value and threshold", "%.6f" % geomValues[3], "%.6f" % geomTargets[3]])
                    OPT_param_print = True  # to prevent repetition of data from OPT and FREQ
                if software == "GAMESS":
                    # in Hartrees per Bohr
                    t.append(["Max Force value and threshold", geomValues[0], geomTargets[0]])
                    t.append(["RMS Force value and threshold", geomValues[1], geomTargets[1]])
            except:
                pass
        # FREQ calculation parameters :
        if 'FREQ' in job_types[i] or ('FREQ' in job_types[i] and 'OPT' in job_types[i]) or ('FREQ' in job_types[i] and 'OPT' in job_types[i] and 'TD' in job_types[i]):
            k = 0
            j = str(k + 1)
            t.append(["Job type: Frequency and thermochemical analysis", " ", " "])
            try:
                t.append(['Temperature', "%.2f K" % json_list[i]["comp_details"]["freq"]["temperature"], "  "])
            except:
                pass
            T_len = False
            try:
                len(json_list[i]["comp_details"]["freq"]["temperature"])
            except KeyError:
                json_list[i]["comp_details"]["freq"]["temperature"] = []
            except TypeError:
                T_len = True
                if T_len is True:
                    try:
                        t.append(['Anharmonic effects', json_list[i]["comp_details"]["freq"]["anharmonicity"], "  "])
                    except KeyError:
                        pass
            if (json_list[i]["comp_details"]["freq"]["temperature"]) != []:
                try:
                    t.append(['Anharmonic effects', json_list[i]["comp_details"]["freq"]["anharmonicity"], "  "])
                except KeyError:
                    pass
                    # TD calculation parameters :
        if 'TD' in job_types[i] or ('FREQ' in job_types[i] and 'OPT' in job_types[i] and 'TD' in job_types[i]):
            k = 0
            j = str(k + 1)
            t.append(["Job type: Time-dependent calculation", " ", " "])
            try:
                t.append(['Number of calculated excited states and spin state', json_list[i]["comp_details"]["excited_states"]["nb_et_states"], np.unique(json_list[i]["results"]["excited_states"]["et_sym"])])
            except KeyError:
                pass
        t.append([" ", " ", " "])
    table = create_table(doc, t)
    set_all_cell_borders(table, False)
    display_vertical_lines(table)
    set_vertical_line(table, 1, False)
    doc.add_paragraph("\nJob type: Geometry optimization")
    table.columns[0].width = Pt(300)
    table.columns[1].width = Pt(80)
    table.columns[2].width = Pt(100)


    ### section 3 : results
    t = []
    add_section_title(doc, "3. RESULTS")
    # Common results / wavefunction :
    t.append(['Total molecular energy', "%.5f hartrees" %
                       data_ref["results"]["wavefunction"]["total_molecular_energy"], " "])
    homo_ind = data_ref["results"]["wavefunction"]["homo_indexes"]
    MO_energies = data_ref["results"]["wavefunction"]["MO_energies"]
    if len(homo_ind) == 2:
        # Unrestricted calculation: two columns of MO energies
        t.append(['Unrestricted calculation', 'Alpha spin MO',
                           'Beta spin MO'])  # indices begin at 0, remove brackets
        t.append(['HOMO number', homo_ind[0] + 1, homo_ind[1] + 1])
        t.append(['LUMO+1 energies', "%.2f eV" % MO_energies[0][homo_ind[0] + 2],
                           "%.2f eV" % MO_energies[1][homo_ind[1] + 2]])
        t.append(['LUMO   energies', "%.2f eV" % MO_energies[0][homo_ind[0] + 1],
                           "%.2f eV" % MO_energies[1][homo_ind[1] + 1]])
        t.append(['HOMO   energies', "%.2f eV" % MO_energies[0][homo_ind[0]],
                           "%.2f eV" % MO_energies[1][homo_ind[1]]])
        t.append(['HOMO-1 energies', "%.2f eV" % MO_energies[0][homo_ind[0] - 1],
                           "%.2f eV" % MO_energies[1][homo_ind[1] - 1]])
    else:
        t.append(['HOMO number', homo_ind[0] + 1, " "])
        t.append(['LUMO+1 energies', "%.2f eV" % MO_energies[0][homo_ind[0] + 2], " "])
        t.append(['LUMO   energies', "%.2f eV" % MO_energies[0][homo_ind[0] + 1], " "])
        t.append(['HOMO   energies', "%.2f eV" % MO_energies[0][homo_ind[0]], " "])
        t.append(['HOMO-1 energies', "%.2f eV" % MO_energies[0][homo_ind[0] - 1], " "])

        # CDFT Indices  table only in full report
    if report_type == 'full':
        t.append([" ", " ", " "])
        try:
            t.append(['CDFT indices: Electron Affinity', "%.4f hartrees" % data_ref["results"]["wavefunction"]["A"],""])
        except KeyError:
            pass
        try:
            t.append(['CDFT indices: Ionisation Potential', "%.4f hartrees" % data_ref["results"]["wavefunction"]["I"], ""])
        except KeyError:
            pass
        try:
            t.append(['CDFT indices: Electronegativity', "%.4f hartrees" % data_ref["results"]["wavefunction"]["Khi"], ""])
        except KeyError:
            pass
        try:
            t.append(['CDFT indices: Hardness', "%.4f hartrees" % data_ref["results"]["wavefunction"]["Eta"], ""])
        except KeyError:
            pass
        try:
            t.append(['CDFT indices: Electrophilicity', "%.4f " % data_ref["results"]["wavefunction"]["Omega"], ""])
        except KeyError:
            pass
        try:
            t.append(['CDFT indices: Electron-flow', "%.4f e-" % data_ref["results"]["wavefunction"]["DeltaN"], ""])
        except KeyError:
            pass
        # Specific calculations results:
    OPT_res_print = False
    for i, jsonfile in enumerate(json_list):
        # OPT calculation results:
        if 'OPT' in job_types[i] or ('FREQ' in job_types[i] and 'OPT' in job_types[i]) \
            or ('FREQ' in job_types[i] and 'OPT' in job_types[i] and 'TD' in job_types[i]) and (OPT_res_print == False):
            j = str(i + 1)
            OPT_res_print = True  # to prevent repetition from OPT and FREQ
            t.append([" ", " ", " "])
            t.append(["Geometry optimization specific results", " ", " "])
            t.append(['Converged nuclear repulsion energy',
                               "%.5f Hartrees" % json_list[i]["results"]["geometry"][
                                   "nuclear_repulsion_energy_from_xyz"], " "])

        # FREQ calculation results:
        if 'FREQ' in job_types[i]  or ('FREQ' in job_types[i] and 'OPT' in job_types[i]) or ('FREQ' in job_types[i] and 'OPT' in job_types[i] and 'TD' in job_types[i]):
            k = 0
            j = str(k + 1)
            t.append([" ", " ", " "])
            t.append(["Frequency and Thermochemistry specific results", " ", " "])
            try:
                rtemper = json_list[i]["comp_details"]["freq"]["temperature"]
            except KeyError:
                rtemper = []
            # ND-arrays
            try:
                vibrational_int = np.array(json_list[i]["results"]["freq"]["vibrational_int"])
            except KeyError:
                vibrational_int = []
            try:
                vibrational_freq = np.array(json_list[i]["results"]["freq"]["vibrational_freq"])
            except KeyError:
                vibrational_freq = []

            if len(vibrational_int) == 0:
                vibrational_int = []
            else:
                # Print number of negative frequencies
                nb_negatives = np.sum(vibrational_freq < 0, axis=0)

            if (len(vibrational_int) != 0) and (rtemper != "N/A"):
                if "zero_point_energy" in json_list[i]["results"]["freq"]:
                    t.append(['Sum of electronic and zero-point energy',
                                       "%.5f Hartrees" % json_list[i]["results"]["freq"]["zero_point_energy"],
                                       " "])
                if "electronic_thermal_energy" in json_list[i]["results"]["freq"]:
                    t.append(["Sum of electronic and thermal energies at  %.2f K" % rtemper,
                                       "%.5f Hartrees" % json_list[i]["results"]["freq"][
                                           "electronic_thermal_energy"], " "])
                if "enthalpy" in json_list[i]["results"]["freq"]:
                    t.append(["Enthalpy at %.2f K" % rtemper,
                                       "%.5f Hartrees" % json_list[i]["results"]["freq"]["enthalpy"], " "])
                if "free_energy" in json_list[i]["results"]["freq"]:
                    t.append(["Gibbs free energy at %.2f K" % rtemper,
                                       "%.5f Hartrees" % json_list[i]["results"]["freq"]["free_energy"], " "])
                if "entropy" in json_list[i]["results"]["freq"]:
                    t.append(["Entropy at %.2f K" % rtemper,
                                       "%.5f Hartrees" % json_list[i]["results"]["freq"]["entropy"], " "])
        # End of the big common result table.

        ## List of tables that are not job associated but dependent of data_ref.
        # Population analysis tables and Fukui condensed values table only in full reports
    if report_type == 'full':
        # Mulliken partial charges table
        try:
            mulliken = data_ref["results"]["wavefunction"]["Mulliken_partial_charges"]
        except KeyError:
            mulliken = []
        # test other population analysis
        try:
            hirsh = data_ref["results"]["wavefunction"]["Hirshfeld_partial_charges"]
        except KeyError:
            hirsh = []
        try:
            cm5 = data_ref["results"]["wavefunction"]["CM5_partial_charges"]
        except KeyError:
            cm5 = []
        if len(mulliken) != 0 :
            # only Mulliken analysis
            mulliken = np.array(mulliken)
            mean_m = np.mean(mulliken)
            dev_m = np.std(mulliken)
            thres_max = mean_m + dev_m
            thres_min = mean_m - dev_m
            if (len(hirsh) == 0) and (len(cm5) == 0):
                ind = np.argsort(mulliken)
                t.append([" ", " " , " "])
                t.append(['Mean Mulliken atomic charge and standard deviation', "%.4f e-" % mean_m , "%.4f e-" % dev_m  ])
                t.append(['Atoms with negatives charges under the standard deviation', "NÂ°" , "Mulliken charge"  ])
                for ielt in ind :
                    if (mulliken[ielt] < thres_min) :
                        t.append([ " " , "%s %d" %(PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]], (1+ielt)), "  %.3f" % mulliken[ielt] ])
                t.append(['Atoms with positives charges over the standard deviation', "NÂ°" , "Mulliken charge"  ])
                for ielt in ind :
                    if (mulliken[ielt] > thres_max) :
                        t.append([ " " , "%s %d" %(PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]], (1+ielt)), "  %+.3f" % mulliken[ielt] ])

            elif (len(hirsh) != 0) and (len(cm5) != 0):
                pass #TODO
                """
                # Hirshfeld and CM5 partial charges table
                cm5 = np.array(cm5)
                mean_h = np.mean(cm5)
                dev_h = np.std(cm5)
                thres_max = mean_h + dev_h
                thres_min = mean_h - dev_h
                ind = np.argsort(cm5)
                res_table.add_row([" ", " ", " "])
                doc.append(NoEscape(r'\begin{center}'))
                with doc.create(Tabular('rlrrr')) as tableau:
                    row_cells = [MultiColumn(5, align='c',
                                             data="Table. Atomic charges population analysis. Selection of the most charged atoms based on Hirshfeld analysis")]
                    tableau.add_row(row_cells)
                    tableau.add_row(["", "Atom and NÂ°", "Hirshfeld charge", "CM5 charge", "Mulliken charge"])
                    tableau.add_hline()
                    for ielt in ind:
                        if (cm5[ielt] < thres_min):
                            tableau.add_row(["",
                                             "%s %d" % (PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]],(1 + ielt)),
                                             "  %+.3f" % cm5[ielt],
                                             "  %+.3f" % hirsh[ielt],
                                             "  %+.3f" % mulliken[ielt]])
                    for ielt in ind:
                        if (cm5[ielt] > thres_max):
                            tableau.add_row(["",
                                             "%s %d" % (PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]],(1 + ielt)),
                                             "  %+.3f" % cm5[ielt],
                                             "  %+.3f" % hirsh[ielt],
                                             "  %+.3f" % mulliken[ielt]])
                    tableau.add_hline()
                doc.append(NoEscape(r'\end{center}'))
                """
        # Fukui condensed values table
        try:
            fplus_lambda_mulliken = data_ref["results"]["wavefunction"]["fplus_lambda_mulliken"]
        except KeyError:
            fplus_lambda_mulliken = []
        try:
            fminus_lambda_mulliken = data_ref["results"]["wavefunction"]["fminus_lambda_mulliken"]
        except KeyError:
            fminus_lambda_mulliken = []
        try:
            fdual_lambda_mulliken = data_ref["results"]["wavefunction"]["fdual_lambda_mulliken"]
        except KeyError:
            fdual_lambda_mulliken = []
        # Fukui condensed values table, check for Hirshfeld values. If present, use them instead.
        try:
            fplus_lambda_hirshfeld = data_ref["results"]["wavefunction"]["fplus_lambda_hirshfeld"]
        except KeyError:
            fplus_lambda_hirshfeld = []
        try:
            fminus_lambda_hirshfeld = data_ref["results"]["wavefunction"]["fminus_lambda_hirshfeld"]
        except KeyError:
            fminus_lambda_hirshfeld = []
        try:
            fdual_lambda_hirshfeld = data_ref["results"]["wavefunction"]["fdual_lambda_hirshfeld"]
        except KeyError:
            fdual_lambda_hirshfeld = []
        # Test if Fdual is available. If not Table is not mandatory.
        if len(fdual_lambda_hirshfeld) > 0:
            pass #TODO
            """
            fdual_lambda_hirshfeld = np.array(fdual_lambda_hirshfeld)
            mean_fd = np.mean(fdual_lambda_hirshfeld)
            dev_fd = np.std(fdual_lambda_hirshfeld)
            thres_max = mean_fd + dev_fd
            thres_min = mean_fd - dev_fd
            ind = np.argsort(fdual_lambda_hirshfeld)
            doc.append(NoEscape(r'\begin{center}'))
            with doc.create(Tabular('p{0cm}rrrrp{0cm}')) as tableau:
                row_cells = [MultiColumn(6, align='c',
                                         data="Table. Selection of the most important condensed Fukui functions based on Hirshfeld charges. ")]
                tableau.add_row(row_cells)
                tableau.add_row(["", "Atom", "atomic dual descriptor (f+ - f-)", "atomic electrophilicity f+",
                                 "atomic nucleophilicity f-", ""])
                tableau.add_hline()
                for ielt in ind:
                    if (fdual_lambda_hirshfeld[ielt] < thres_min):
                        tableau.add_row(["", "%s %d" % (
                        PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]], (1 + ielt)),
                                         "  %.2f" % fdual_lambda_hirshfeld[ielt],
                                         "  %.2f" % fplus_lambda_hirshfeld[ielt],
                                         "  %.2f" % fminus_lambda_hirshfeld[ielt],
                                         ""])
                for ielt in ind:
                    if (fdual_lambda_hirshfeld[ielt] > thres_max):
                        tableau.add_row(["", "%s %d" % (
                        PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]], (1 + ielt)),
                                         "  %.2f" % fdual_lambda_hirshfeld[ielt],
                                         "  %.2f" % fplus_lambda_hirshfeld[ielt],
                                         "  %.2f" % fminus_lambda_hirshfeld[ielt],
                                         ""])
                tableau.add_hline()
            doc.append(NoEscape(r'\end{center}'))
            """
        # Test if Fdual Hirshfeld is not available check for Mulliken one. If not Table is not mandatory.
        elif len(fdual_lambda_mulliken) > 0:
            pass #TODO
            """
            fdual_lambda_mulliken = np.array(fdual_lambda_mulliken)
            mean_fd = np.mean(fdual_lambda_mulliken)
            dev_fd = np.std(fdual_lambda_mulliken)
            thres_max = mean_fd + dev_fd
            thres_min = mean_fd - dev_fd
            ind = np.argsort(fdual_lambda_mulliken)
            doc.append(NoEscape(r'\begin{center}'))
            with doc.create(Tabular('p{0cm}rrrrp{0cm}')) as tableau:
                row_cells = [MultiColumn(6, align='c',
                                         data="Table. Selection of the most important condensed Fukui functions based on Mulliken charges. ")]
                tableau.add_row(row_cells)
                tableau.add_row(["", "Atom", "atomic dual descriptor (f+ - f-)", "atomic electrophilicity f+",
                                 "atomic nucleophilicity f-", ""])
                tableau.add_hline()
                for ielt in ind:
                    if (fdual_lambda_mulliken[ielt] < thres_min):
                        tableau.add_row(["", "%s %d" % (
                        PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]], (1 + ielt)),
                                         "  %.2f" % fdual_lambda_mulliken[ielt],
                                         "  %.2f" % fplus_lambda_mulliken[ielt],
                                         "  %.2f" % fminus_lambda_mulliken[ielt],
                                         ""])
                for ielt in ind:
                    if (fdual_lambda_mulliken[ielt] > thres_max):
                        tableau.add_row(["", "%s %d" % (
                        PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][ielt]], (1 + ielt)),
                                         "  %.2f" % fdual_lambda_mulliken[ielt],
                                         "  %.2f" % fplus_lambda_mulliken[ielt],
                                         "  %.2f" % fminus_lambda_mulliken[ielt],
                                         ""])
                tableau.add_hline()
            doc.append(NoEscape(r'\end{center}'))
            """
    table = create_table(doc, t)
    set_all_cell_borders(table, False)
    display_vertical_lines(table)
    set_vertical_line(table, 1, False)
    doc.add_paragraph("\nJob type: Geometry optimization")
    table.columns[0].width = Pt(300)
    table.columns[1].width = Pt(80)
    table.columns[2].width = Pt(100)
    ## List of figures. Beware insertion based on files. Should be given through arguments!
    # figure with MO not available in text report type
    if report_type != 'text':
        # Test if calculation is unrestricted (alpha and beta spin electrons)
        if len(homo_ind) == 2:
            # Unrestricted calculation: treat the alpha orbitals first
            nomPng = "temp/img-MO-homo_alpha.png"
            nomPng2 = "temp/img-MO-homo_beta.png"
            nomPng3 = "temp/img-MO-lumo_alpha.png"
            nomPng4 = "temp/img-MO-lumo_beta.png"
            if (os.path.isfile(nomPng)):
                if (os.path.isfile(nomPng2)):
                    figure_two_col(doc, nomPng, nomPng2, "Representation of the HOMO of spin alpha (left) and spin beta (right).")
                    if (os.path.isfile(nomPng3)):
                        if (os.path.isfile(nomPng4)):
                            figure_two_col(doc, nomPng3, nomPng4, "Representation of the LUMO of spin alpha (left) and spin beta (right).")
                elif (os.path.isfile(nomPng3)):
                    figure_two_col(doc, nomPng, nomPng3, "Representation of the Frontier Molecular Orbitals HOMO (left) and LUMO (right) of spin alpha.")
        # For restricted calculation print HOMO and LUMO and use two cameras instead
        else:
            nomPng = "temp/img-MO-homo.png"
            nomPng2 = "temp/img-MO-homo_cam2.png"
            nomPng3 = "temp/img-MO-lumo.png"
            nomPng4 = "temp/img-MO-lumo_cam2.png"
            if (os.path.isfile(nomPng)):
                if (os.path.isfile(nomPng2)):
                    figure_two_col(doc, nomPng, nomPng2, "Representation of the HOMO from two points of view.")
                    if (os.path.isfile(nomPng3)):
                        if (os.path.isfile(nomPng4)):
                            figure_two_col(doc, nomPng3, nomPng4, "Representation of the LUMO from two points of view.")
                elif (os.path.isfile(nomPng3)):
                    figure_two_col(doc, nomPng, nomPng3, "Representation of the Frontier Molecular Orbitals HOMO (left) and LUMO (right).")

    # Figures of that are only printed in full type report
    if report_type == 'full':
        # figures of Fukui functions if calculated
        nomPng = "temp/img-fukui-SP_plus.png"
        nomPng2 = "temp/img-fukui-SP_plus_cam2.png"
        nomPng3 = "temp/img-fukui-SP_minus.png"
        nomPng4 = "temp/img-fukui-SP_minus_cam2.png"
        if (os.path.isfile(nomPng)):
            if (os.path.isfile(nomPng2)):
                figure_two_col(doc, nomPng, nomPng2, "Representation of the F+ function from two points of view. The Blue color indicate the most electrophilic regions.")
                if (os.path.isfile(nomPng3)):
                    if (os.path.isfile(nomPng4)):
                        figure_two_col(doc, nomPng3, nomPng4, "Representation of the F- function from two points of view. The Blue color indicate the most nucleophilic regions.")
            elif (os.path.isfile(nomPng3)):
                figure_two_col(doc, nomPng, nomPng3, "Representation of the electrophilic (left) and nucleophilic (right) fukui functions.")
        nomPng = "temp/img-Fdual.png"
        nomPng2 = "temp/img-Fdual_cam2.png"
        if (os.path.isfile(nomPng)):
            if (os.path.isfile(nomPng2)):
                figure_two_col(doc, nomPng, nomPng2, "Representation of the Dual descriptor from two points of view. Electrophilic and nucleophilic regions correspond to blue and white surfaces.")
            elif (os.path.isfile(nomPng3)):
                figure_two_col(doc, nomPng, nomPng2, "Representation of the Dual descriptor. Electrophilic and nucleophilic regions correspond to blue and white surfaces.")

        # figure with ESP
        nomPng = "temp/img-MEP_fixed.png"
        nomPng2 = "temp/img-MEP.png"
        if (os.path.isfile(nomPng)):
            if (os.path.isfile(nomPng2)):
                figure_two_col(doc, nomPng, nomPng2, "Representations of the Molecular Electrostatic Potential mapped on the electron density (cutoff value of 0.002 e-/bohr3). On the left, red, blue and green regions correspond to negative values < -0.06 a.u., positive values > 0.08 a.u. and neutral values respectively. On the right, the scale is set automatically to highlight the minimum values in red and the maximum values in blues.")
            elif (not os.path.isfile(nomPng2)):
                figure_one_col(doc, nomPng, "Representations of the Molecular Electrostatic Potential mapped on the electron density (cutoff value of 0.002 e-/bohr3). On the left, red, blue and green regions correspond to negative values < -0.06 a.u., positive values > 0.08 a.u. and neutral values respectively.")

        # External picture generated by AIMAll
        nomPng = "temp/img-AIM-BCP-rho.png"
        if (not os.path.isfile(nomPng)):
            print(nomPng + " not found. It can not be added to the report.\n")
        else:
            figure_one_col(doc, nomPng, "Electron density value at each of the bond critical points calculated with the AIMAll program.")

    # Specific OPT and FREQ report tables
    for i, jsonfile in enumerate(json_list):
        # Normal modes table
        if 'FREQ' in job_types[i] or ('FREQ' in job_types[i] and 'OPT' in job_types[i]) or ('FREQ' in job_types[i] and 'OPT' in job_types[i] and 'TD' in job_types[i]):
            k = 0
            try:
                vibrational_int = np.array(json_list[i]["results"]["freq"]["vibrational_int"])
            except KeyError:
                vibrational_int = []
            try:
                vibrational_freq = np.array(json_list[i]["results"]["freq"]["vibrational_freq"])
            except KeyError:
                vibrational_freq = []
            try:
                vibrational_sym = np.array(json_list[i]["results"]["freq"]["vibrational_sym"])
            except KeyError:
                vibrational_sym = np.array(["N/A" for _ in range(len(vibrational_int))])
            # filtering & orderering
            if len(vibrational_int) == 0:
                vibrational_int = []
            else:
                vib_filter = vibrational_int > 50.
                vib_order = np.argsort(vibrational_freq[vib_filter])[::-1]
                vibrational_int = vibrational_int[vib_filter][vib_order]
                vibrational_freq = vibrational_freq[vib_filter][vib_order]
                vibrational_sym = vibrational_sym[vib_filter][vib_order]
            t = []
            t.append(["Table. Most intense (> 50 km/mol) molecular vibrations in wavenumbers", "", ""])
            t.append(["Frequencies", "Intensity", "Symmetry"])
            for k in range(len(vibrational_freq)):
                t.append(["%d" % vibrational_freq[k], "%d" % vibrational_int[k], vibrational_sym[k]])
            t.append(["", "", ""])
            table = create_table(doc, t)
            table.rows[0].cells[0].merge(table.rows[0].cells[2])
            set_horizontal_line(table, 2, True, 8)
            set_horizontal_line(table, 3, True, 8)

        # TD calculation results :
        if 'TD' in job_types[i] or ('FREQ' in job_types[i] and 'OPT' in job_types[i] and 'TD' in job_types[i]):
            j = str(i + 1)
            try:
                et_energies = json_list[i]["results"]["excited_states"]["et_energies"]
            except KeyError:
                et_energies = []
            # Conversion of wavenumbers to nm
            et_nm = [nm_to_wnb / x for x in et_energies]
            rnbExci = len(et_energies)
            if rnbExci != 0 and et_energies != 'N/A':
                # 2 kind of table including the charge transfer data when discretization is done (SI and Full)
                if report_type == 'text':
                    t = []
                    if rnbExci < 21:
                        t.append(["Table. Results concerning the calculated mono-electronic excitations."])
                        td_selection = 0
                    else:  # select the excited states when there is more than 20 calculated excited states.
                        t.append(["Table. First five calculated mono-electronic excitations and those with f > 0.1 or R > 10."])
                        td_selection = 1
                    t.append(["E.S.", "Symmetry", " nm ", "", "", "R", "Excitation description : initial OM - ending OM (% if > 5%)"])
                    t.append(["", "", "", "", "", "", ""])
                    for j in range(rnbExci):
                        try:
                            etr_i = json_list[i]["results"]["excited_states"]["et_rot"][j]
                        except KeyError:
                            etr_i = 0.
                        trans = json_list[i]["results"]["excited_states"]["et_transitions"][j]
                        # Form the string description of the escitation. Based on MO init -> MO end (%coeff)
                        CIS = " "
                        for subtrans, ST in enumerate(trans):
                            coeff = int(ST[2] ** 2 * 100)
                            if coeff > 5.:
                                if len(homo_ind) == 2:  # Unrestricted calculation Spin needs to be printed
                                    if ST[0][1] == 0:
                                        spin_init = "a"  # spin alpha to print in the table
                                    elif ST[0][1] == 1:
                                        spin_init = "b"  # spin beta to print in the table
                                    else:
                                        spin_init = ""  # unrecognised spin
                                    if ST[1][1] == 0:
                                        spin_end = "a"  # spin alpha to print in the table
                                    elif ST[1][1] == 1:
                                        spin_end = "b"  # spin beta to print in the table
                                    else:
                                        spin_end = ""  # unrecognised spin
                                    CIS += str(ST[0][0] + 1) + spin_init + "-" + str(
                                        ST[1][0] + 1) + spin_end + " (" + str(coeff) + ") "
                                if len(homo_ind) == 1:  # Restricted calculation Spin is omitted
                                    CIS += str(ST[0][0] + 1) + "-" + str(ST[1][0] + 1) + "(" + str(
                                        coeff) + "); "
                        # select the excited states when there is more than 20 calculated excited states. based on oscillator strength and rotational strength
                        if (td_selection == 0) or ((td_selection == 1) and ((j < 5) or \
                                                                            (json_list[i]["results"][
                                                                                 "excited_states"]["et_oscs"][
                                                                                 j] > 0.1) or \
                                                                            ((etr_i == "N/A") or (
                                                                                    abs(etr_i) > 10.)))):
                            t.append([(1 + j),
                                             json_list[i]["results"]["excited_states"]["et_sym"][j],
                                             "%d " % et_nm[j],
                                             "%d " % et_energies[j],
                                             "%.3f" % json_list[i]["results"]["excited_states"]["et_oscs"][j],
                                             "%.1f" % etr_i,
                                             # Printing only transitions over 5%
                                             "%s" % CIS
                                             ])
                    t.append(["", "", "", "", "", "", ""])
                else:
                    if rnbExci < 21:
                        t.append(["Table. Results concerning the calculated mono-electronic excitations."])
                        td_selection = 0
                    else:  # select the excited states when there is more than 20 calculated excited states.
                        t.append(["Table. First five calculated mono-electronic excitations and those with f > 0.1 or R > 10."])
                        td_selection = 1
                    t.append(["E.S.", "Symmetry", " nm ", "","", "R", "", "","", "Excitation description : initial OM - ending OM (% if > 5%)"])
                    t.append(["", "", "", "","", "", "", "","", ""])
                    for j in range(rnbExci):
                        try:
                            etr_i = json_list[i]["results"]["excited_states"]["et_rot"][j]
                        except KeyError:
                            etr_i = 0.
                        trans = json_list[i]["results"]["excited_states"]["et_transitions"][j]
                        # Form the string description of the escitation. Based on MO init -> MO end (%coeff)
                        CIS = " "
                        for subtrans, ST in enumerate(trans):
                            coeff = int(ST[2] ** 2 * 100)
                            if coeff > 5.:
                                if len(homo_ind) == 2:  # Unrestricted calculation Spin needs to be printed
                                    if ST[0][1] == 0:
                                        spin_init = "a"  # spin alpha to print in the table
                                    elif ST[0][1] == 1:
                                        spin_init = "b"  # spin beta to print in the table
                                    else:
                                        spin_init = ""  # unrecognised spin
                                    if ST[1][1] == 0:
                                        spin_end = "a"  # spin alpha to print in the table
                                    elif ST[1][1] == 1:
                                        spin_end = "b"  # spin beta to print in the table
                                    else:
                                        spin_end = ""  # unrecognised spin
                                    CIS += str(ST[0][0] + 1) + spin_init + "-" + str(
                                        ST[1][0] + 1) + spin_end + " (" + str(coeff) + ") "
                                if len(homo_ind) == 1:  # Restricted calculation Spin is omitted
                                    CIS += str(ST[0][0] + 1) + "-" + str(ST[1][0] + 1) + "(" + str(
                                        coeff) + "); "
                        # select the excited states when there is more than 20 calculated excited states. based on oscillator strength and rotational strength
                        if (td_selection == 0) or ((td_selection == 1) and ((j < 5) or \
                                                                            (json_list[i]["results"][
                                                                                 "excited_states"]["et_oscs"][
                                                                                 j] > 0.1) or \
                                                                            ((etr_i == "N/A") or (
                                                                                    abs(etr_i) > 10.)))):
                            t.append([(1 + j),
                                             json_list[i]["results"]["excited_states"]["et_sym"][j],
                                             "%d " % et_nm[j],
                                             "%d " % et_energies[j],
                                             "%.3f" % json_list[i]["results"]["excited_states"]["et_oscs"][j],
                                             "%.1f" % etr_i,
                                             "%.2f" % json_list[i]["results"]["excited_states"]["Tozer_lambda"][
                                                 j],
                                             "%.2f" % json_list[i]["results"]["excited_states"]["d_ct"][j],
                                             "%.2f" % json_list[i]["results"]["excited_states"]["q_ct"][j],
                                             # Printing only transitions over 5%
                                             "%s" % CIS
                                             ])


                create_table(doc, t)
    # UV visible Absorption and Circular dischroism plots
    nomPng3 = "temp/img-UV-Abso-Spectrum.png"
    if (not os.path.isfile(nomPng3)):
        print("No PNG named " + nomPng3 + " found. The spectrum can not be added to the report.\n")
    else:
        figure_one_col(doc, nomPng3, "Calculated UV visible Absorption spectrum with a gaussian broadening (FWHM = 3000 cm-1)")

    nomPng4 = "temp/img-UV-CD-Spectrum.png"
    if (not os.path.isfile(nomPng4)):
        print("No PNG named " + nomPng4 + " found. The spectrum can not be added to the report.\n")
    else:
        figure_one_col(doc, nomPng4, "Calculated Circular Dichroism spectrum with a gaussian broadening (FWHM = 3000 cm-1)")

        # figures not available in text report type
        if report_type != 'text':

            # figure with EDD
            nomPng = "temp/img-EDD-S1.png"
            nomPng_cam2 = "temp/img-EDD-S1_cam2.png"
            nomPng2 = "temp/img-EDD-S2.png"
            nomPng2_cam2 = "temp/img-EDD-S2_cam2.png"
            if (os.path.isfile(nomPng)):
                if (os.path.isfile(nomPng_cam2)):
                    figure_two_col(doc, nomPng, nomPng_cam2, "Representation of the Electron Density Difference (S1-S0) from two points of view.")
                    if (os.path.isfile(nomPng2)):
                        if (os.path.isfile(nomPng2_cam2)):
                            figure_two_col(doc, nomPng2, nomPng2_cam2, "Representation of the Electron Density Difference (S2-S0) from two points of view.")
                elif (os.path.isfile(nomPng2)):
                    figure_two_col(doc, nomPng, nomPng2, "Representation of the Electron Density Difference (S1-S0 left) and (S2-S0 right). The excited electron and the hole regions are indicated by respectively blue and white surfaces.")
            else:
                nomPng = "temp/img-EDD-1.png"
                nomPng_cam2 = "temp/img-EDD-1_cam2.png"
                nomPng2 = "temp/img-EDD-2.png"
                nomPng2_cam2 = "temp/img-EDD-2_cam2.png"
                if (os.path.isfile(nomPng)):
                    if (os.path.isfile(nomPng_cam2)):
                        figure_two_col(doc, nomPng, nomPng_cam2, "Representation of the Electron Density Difference (ES1-GS) from two points of view.")
                        if (os.path.isfile(nomPng2)):
                            if (os.path.isfile(nomPng2_cam2)):
                                figure_two_col(doc, nomPng2, nomPng2_cam2, "Representation of the Electron Density Difference (ES2-GS) from two points of view.")
                    elif (os.path.isfile(nomPng2)):
                        figure_two_col(doc, nomPng, nomPng2, "Representation of the Electron Density Difference (ES1-GS left) and (ES2-GS right). The excited electron and the hole regions are indicated by respectively blue and white surfaces.")

        # Specific OPT_ES report tables
        #TODO
        """
        for i, jsonfile in enumerate(json_list):
            # TD emission calculation results :
            if 'OPT_ES' in job_types[i]:
                j = str(i + 1)
                # res_table.add_row(["Optimization of Time-dependent excited state specific results", " " , " "])
                emi_state = [int(s) for s in json_list[i]["metadata"]["log_file"] if s.isdigit()][0]
                emi_index = emi_state - 1
                try:
                    emi_energy = json_list[i]["results"]["excited_states"]["et_energies"][emi_index]
                except KeyError:
                    emi_energy = 0.0
                # Conversion of wavenumbers to nm
                et_nm = [nm_to_wnb / emi_energy]
                rnbExci = len([emi_energy])
                if rnbExci != 0 and emi_energy != 0.0:
                    doc.append(NoEscape(r'\begin{center}'))
                    # 2 kind of table including the charge transfer data when discretization is done (SI and Full)
                    if report_type == 'text':
                        with doc.create(Tabular('rrrrrrp{6cm}')) as tableau:
                            row_cells = [MultiColumn(7, align='c',
                                                     data="Table. Results concerning the calculated mono-electronic optimization excitation")]
                            tableau.add_row(row_cells)
                            tableau.add_row(["E.S.", "Symmetry", " nm ", NoEscape(r"cm$^{-1}$"),
                                             italic("f"), "R", "Excitation description in %"])
                            tableau.add_hline()
                            try:
                                etr_i = json_list[i]["results"]["excited_states"]["et_rot"][emi_index]
                            except KeyError:
                                etr_i = "N/A"
                            trans = json_list[i]["results"]["excited_states"]["et_transitions"][emi_index]
                            # Form the string description of the escitation. Based on MO init -> MO end (%coeff)
                            CIS = " "
                            for subtrans, ST in enumerate(trans):
                                coeff = int(ST[2] ** 2 * 100)
                                if len(homo_ind) == 2:  # Unrestricted calculation Spin needs to be printed
                                    if ST[0][1] == 0:
                                        spin_init = "a"  # spin alpha to print in the table
                                    elif ST[0][1] == 1:
                                        spin_init = "b"  # spin beta to print in the table
                                    else:
                                        spin_init = ""  # unrecognised spin
                                    if ST[1][1] == 0:
                                        spin_end = "a"  # spin alpha to print in the table
                                    elif ST[1][1] == 1:
                                        spin_end = "b"  # spin beta to print in the table
                                    else:
                                        spin_end = ""  # unrecognised spin
                                    CIS += str(ST[0][0] + 1) + spin_init + "-" + str(
                                        ST[1][0] + 1) + spin_end + " (" + str(coeff) + ") "
                                if len(homo_ind) == 1:  # Restricted calculation Spin is omitted
                                    CIS += str(ST[0][0] + 1) + "->" + str(ST[1][0] + 1) + " (" + str(coeff) + ") "
                            tableau.add_row([(2),
                                             json_list[i]["results"]["excited_states"]["et_sym"][emi_index],
                                             "%d " % et_nm[emi_index],
                                             "%d " % emi_energy,
                                             "%.3f" % json_list[i]["results"]["excited_states"]["et_oscs"][
                                                 emi_index],
                                             "%s" % etr_i,  # "%.3f" % etr_i, ToDo NOT working when et rot  = N/A
                                             # trying to reduce CIS size size it can too large too fit in the page
                                             "%s" % CIS
                                             ])
                            tableau.add_hline()
                    else:
                        with doc.create(Tabular('rrrrrrrrrp{6cm}')) as tableau:
                            row_cells = [MultiColumn(10, align='c',
                                                     data="Table. Results concerning the calculated mono-electronic optimization excitation")]
                            tableau.add_row(row_cells)
                            tableau.add_row(["E.S.", "Symmetry", " nm ", NoEscape(r"cm$^{-1}$"),
                                             italic("f"), "R", NoEscape(r"$\Lambda$"), NoEscape(r"d$_{CT}$"),
                                             NoEscape(r"q$_{CT}$"), "Excitation description in %"])
                            tableau.add_hline()
                            try:
                                etr_i = json_list[i]["results"]["excited_states"]["et_rot"][emi_index]
                            except KeyError:
                                etr_i = "N/A"
                            trans = json_list[i]["results"]["excited_states"]["et_transitions"][emi_index]
                            # Form the string description of the escitation. Based on MO init -> MO end (%coeff)
                            CIS = " "
                            for subtrans, ST in enumerate(trans):
                                coeff = int(ST[2] ** 2 * 100)
                                if len(homo_ind) == 2:  # Unrestricted calculation Spin needs to be printed
                                    if ST[0][1] == 0:
                                        spin_init = "a"  # spin alpha to print in the table
                                    elif ST[0][1] == 1:
                                        spin_init = "b"  # spin beta to print in the table
                                    else:
                                        spin_init = ""  # unrecognised spin
                                    if ST[1][1] == 0:
                                        spin_end = "a"  # spin alpha to print in the table
                                    elif ST[1][1] == 1:
                                        spin_end = "b"  # spin beta to print in the table
                                    else:
                                        spin_end = ""  # unrecognised spin
                                    CIS += str(ST[0][0] + 1) + spin_init + "-" + str(
                                        ST[1][0] + 1) + spin_end + " (" + str(coeff) + ") "
                                if len(homo_ind) == 1:  # Restricted calculation Spin is omitted
                                    CIS += str(ST[0][0] + 1) + "->" + str(ST[1][0] + 1) + " (" + str(coeff) + ") "
                            tableau.add_row([(2),
                                             json_list[i]["results"]["excited_states"]["et_sym"][emi_index],
                                             "%d " % et_nm[emi_index],
                                             "%d " % emi_energy,
                                             "%.3f" % json_list[i]["results"]["excited_states"]["et_oscs"][
                                                 emi_index],
                                             "%s" % etr_i,  # "%.3f" % etr_i, ToDo NOT working when et rot  = N/A
                                             "%.2f" % json_list[i]["results"]["excited_states"]["Tozer_lambda"][
                                                 emi_index],
                                             "%.2f" % json_list[i]["results"]["excited_states"]["d_ct"][emi_index],
                                             "%.2f" % json_list[i]["results"]["excited_states"]["q_ct"][emi_index],
                                             # trying to reduce CIS size size it can too large too fit in the page
                                             "%s" % CIS
                                             ])
                            tableau.add_hline()
                    doc.append(NoEscape(r'\end{center}'))
"""
                    # UV visible Emission and Circular dischroism plots
        nomPng3 = "temp/img-UV-Emi-Spectrum.png"
        if (not os.path.isfile(nomPng3)):
            print("No PNG named " + nomPng3 + " found. The spectrum can not be added to the report.\n")
        else:
            figure_one_col(doc, nomPng3, "Calculated UV visible Emission spectrum with a gaussian broadening (FWHM = 3000 cm-1)")

        nomPng4 = "temp/img-UV-CD-Emi-Spectrum.png"
        if (not os.path.isfile(nomPng4)):
            print("No PNG named " + nomPng4 + " found. The spectrum can not be added to the report.\n")
        else:
            figure_one_col(doc, nomPng4, "Calculated Circular Dichroism Emission spectrum with a gaussian broadening (FWHM = 3000 cm-1)")

        if report_type != 'text':
            # figure with EDD for emission
            nomPng = "temp/img-emi-EDD-S1.png"
            if (os.path.isfile(nomPng)):
                figure_one_col(doc, nomPng, "Representation of the Electron Density Difference (S1-GS) after optimization of the excited state. The excited electron and the hole regions are indicated by respectively white and blue surfaces to ease comparison with the corresponding absorption transition.")

        ######## Table of atomic coordinates
        # doc.append(NoEscape(r'\clearpage'))
        atomic = data_ref["results"]["geometry"]["elements_3D_coords_converged"]
        oxtrf, oytrf, oxtrd, oytrd = ('N/A', 'N/A', 'N/A', 'N/A')

        t = []
        t.append(["", "Table. Converged cartesian atomic coordinates in Angstroms", "", "", "", ""])
        t.append(["", "Atom", 'X', 'Y', 'Z'])
        atoms = np.array(json_list[i]["results"]["geometry"]["elements_3D_coords_converged"]).reshape((-1, 3))
        for j, a in enumerate(atoms):
            t.append(["", PeriodicTable().element[json_list[i]['molecule']["atoms_Z"][j]], "%.4f" % a[0], "%.4f" % a[1], "%.4f" % a[2], ""])
        table = create_table(doc, t)
        set_horizontal_line(table, 1, True, 8)
        set_horizontal_line(table, len(table.rows), True, 8)
        table.rows[0].cells[1].merge(table.rows[0].cells[4])
    doc.add_paragraph()
    add_footer(doc)
    doc.save(dirname + "_" + report_type + "_report.docx")
